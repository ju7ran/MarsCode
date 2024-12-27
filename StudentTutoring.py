from openai import OpenAI
from dotenv import load_dotenv
import os
import streamlit as st
from docx import Document

load_dotenv()
from docx2pdf import convert
import pythoncom

# 设置 OpenAI 的 API 密钥（如果你使用的是 OpenAI 的模型）
os.environ["http_proxy"] = "http://127.0.0.1:7897"
os.environ["https_proxy"] = "http://127.0.0.1:7897"

client = OpenAI()


# 定义一个函数来调用模型
def get_completion(user_input, model="gpt-3.5-turbo"):
    instruction = """
    你是一位专业的大模型辅导老师，为学员提供个性化的学习建议，帮助他们更好地掌握大模型知识和技能。
    """
    # output = """
    # 并以字符串格式输出
    # """

    examples = """
    # 示例1
        Q：您现在在那个城市，是否在职，所从事的工作是什么？
        A：深圳，在职，外贸业务员
        Q：对大模型有多少认知，了解多少原理与技术点？
        A：0基础
        Q：学习大模型的最核心需求是什么？
        A：提高核心竞争力
        Q：是否有python编程基础或者其他编程基础，有没有写过代码？
        A：无
        Q：每天能花多少时间用于学习，大致空闲时间点处于什么时段?
        A：3小时，晚上7点后
        Q：除以上五点外是否还有其他问题想要补充。如有请按照如下格式进行补充

        给学员的回复是
        你现在对大模型有基本的认知，你可以在公众号或者知乎上了解一些大模型的知识，了解大模型的一些技术发展，你的学习时间比较多，前期可以多花点时间看录播，每节课都有课后回放的，这样你在后面的学习中会比较容易跟上老师，大模型的主要语言是Python，你之前有接触编程，Python的预习视频可以快速过一遍，大模型现在工作还是很火热的，大模型的前景发展也是非常好的，现在国内大模型的发展处于刚起步阶段，还是有很多机会的，希望你能在这里学有所成。


        # 示例2
        Q：您现在在那个城市，是否在职，所从事的工作是什么？
        A：北京，在职，农业相关
        Q：对大模型有多少认知，了解多少原理与技术点？
        A：比较浅薄
        Q：学习大模型的最核心需求是什么？
        A：个人能力提升和业务需要
        Q：是否有python编程基础或者其他编程基础，有没有写过代码？
        A：有
        Q：每天能花多少时间用于学习，大致空闲时间点处于什么时段?
        A：3个小时左右，晚上18点以后
        Q：除以上五点外是否还有其他问题想要补充。如有请按照如下格式进行补充

        给学员的回复是
        作为在北京从事农业相关工作的同学，虽然你对大模型的认知程度比较浅，但你
        拥有 Python 编程基础并且写过代码，这对于学习大模型来说是很好的条件，因
        为 Python 是学习大模型的主要语言。推荐你看一下我们提供的预习课程来补充
        一下知识体系。个人能力提升和业务需要符合当前 AI 在农业领域的发展趋势。
        每天在晚上 18 点以后可以安排约 3 个小时的学习时间，这样的时间安排非常充
        裕。凭借你的编程背景和学习投入，转型为 AI 项目管理是可行的，国内现在 AI
        领域虽然处于起步阶段，但随着人工智能技术的快速发展，其应用前景非常广阔，
        现在正是学习并把握行业发展机遇的好时机
    """
    prompt = f"""
        {instruction}
        {examples}
        用户输入：
        {user_input}
        限制：限制
            - 只提供与大模型学习相关的建议，拒绝回答与大模型无关的问题。
            - 所输出的内容必须按照给定的格式进行组织，不能偏离框架要求。
            - 建议内容要具体、可行，具有针对性和可操作性。
    """
    print(prompt)
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,  # 模型输出的随机性，0 表示随机性最小
        n=4
    )
    return response.choices[0].message.content


def get_completion_pdf(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,  # 模型输出的随机性，0 表示随机性最小
        n=4
    )
    return response.choices[0].message.content


# Streamlit 界面设置
st.title("智能学员辅导系统")
st.write("在下方输入您的问题，点击生成按钮，AI将为您生成回答。")

# 获取用户输入的各个问题
city_work = st.text_input("Q1：您现在在那个城市，是否在职，所从事的工作是什么？")
recognition = st.text_input("Q2：对大模型有多少认知，了解多少原理与技术点？")
core_need = st.text_input("Q3：学习大模型的最核心需求是什么？")
programming = st.text_input("Q4：是否有python编程基础或者其他编程基础，有没有写过代码？")
study_time = st.text_input("Q5：每天能花多少时间用于学习，大致空闲时间点处于什么时段?")
other_questions = st.text_input("Q6：除以上五点外是否还有其他问题想要补充。如有请按照如下格式进行补充")

# 输入文本
# user_input = st.text_input("请输入您的问题:")

if st.button("生成"):
    if all([city_work, recognition, core_need, programming, study_time, other_questions]):
        with st.spinner('AI 正在思考...'):
            # 将用户输入的内容组合成一个整体的输入
            user_input = f"""
            Q：您现在在那个城市，是否在职，所从事的工作是什么？
            A：{city_work}
            Q：对大模型有多少认知，了解多少原理与技术点？
            A：{recognition}
            Q：学习大模型的最核心需求是什么？
            A：{core_need}
            Q：是否有python编程基础或者其他编程基础，有没有写过代码？
            A：{programming}
            Q：每天能花多少时间用于学习，大致空闲时间点处于什么时段?
            A：{study_time}
            Q：除以上五点外是否还有其他问题想要补充。如有请按照如下格式进行补充
            A：{other_questions}
            """
            # 调用模型生成回复
            output = get_completion(user_input)

            # 创建并保存为 DOCX 文件
            doc = Document('学习规划.docx')
            for para in doc.paragraphs:
                if '[]' in para.text:  # 替换文档中出现的'old_text'
                    para.text = para.text.replace('[]', output)
            doc.add_paragraph(output)

            modified_file = '智泊AI学习规划.docx'
            doc.save(modified_file)
            st.success(output)

            # 确保 COM 初始化
            pythoncom.CoInitialize()  # 初始化 COM 库

            # 转换 DOCX 为 PDF
            convert(modified_file, f"{modified_file.replace('.docx', '.pdf')}")
            pdf_file = modified_file.replace('.docx', '.pdf')

            # 提供文件下载按钮
            with open(pdf_file, "rb") as f:
                st.download_button(
                    label="点击下载完整学习规划",
                    data=f,
                    file_name=pdf_file,
                    mime="application/pdf"
                )
    else:
        st.error("请确保所有问题都已填写。")

# 添加一个信息栏
st.sidebar.info("这个应用程序使用 OpenAI 的 GPT-3 模型进行交互。")